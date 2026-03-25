"""DOCX 导出器：将 HTML 转换为 Word 文档"""

from pathlib import Path

from bs4 import BeautifulSoup

from exporters.base import BaseExporter
from utils.logger import get_logger

logger = get_logger("docx_exporter")


class DocxExporter(BaseExporter):
    """将 HTML 转换为 DOCX"""

    def export(self, html_content: str, url: str, filepath: str) -> bool:
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("DOCX 导出需要 python-docx，请运行: pip install python-docx")

        Path(filepath).parent.mkdir(parents=True, exist_ok=True)

        try:
            cleaned = self.clean_html(html_content, url)
            soup = BeautifulSoup(cleaned, "html.parser")
            body = soup.find("body") or soup

            doc = Document()

            # 添加来源信息
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"来源: {url}")
            run.font.size = Pt(8)
            run.font.italic = True

            # 遍历 body 的直接子元素
            for elem in body.children:
                if not hasattr(elem, 'name') or elem.name is None:
                    # 纯文本节点
                    text = elem.strip() if isinstance(elem, str) else ""
                    if text:
                        doc.add_paragraph(text)
                    continue

                self._convert_element(doc, elem)

            doc.save(filepath)
            return True
        except Exception as e:
            logger.error(f"DOCX 导出失败: {url} -> {e}")
            return False

    def _convert_element(self, doc, elem):
        """递归转换 HTML 元素为 docx 元素"""
        tag = elem.name.lower() if elem.name else ""

        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(tag[1])
            text = elem.get_text(strip=True)
            if text:
                doc.add_heading(text, level=min(level, 9))

        elif tag == 'p':
            text = elem.get_text(strip=True)
            if text:
                doc.add_paragraph(text)

        elif tag in ('ul', 'ol'):
            for li in elem.find_all('li', recursive=False):
                text = li.get_text(strip=True)
                if text:
                    doc.add_paragraph(text, style='List Bullet' if tag == 'ul' else 'List Number')

        elif tag == 'table':
            self._convert_table(doc, elem)

        elif tag == 'blockquote':
            text = elem.get_text(strip=True)
            if text:
                p = doc.add_paragraph(text)
                p.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else None

        elif tag in ('div', 'section', 'article', 'main'):
            for child in elem.children:
                if hasattr(child, 'name') and child.name:
                    self._convert_element(doc, child)
                elif isinstance(child, str) and child.strip():
                    doc.add_paragraph(child.strip())

        elif tag == 'br':
            pass  # 忽略

        elif tag not in ('script', 'style', 'noscript', 'iframe'):
            text = elem.get_text(strip=True)
            if text and len(text) > 5:
                doc.add_paragraph(text)

    def _convert_table(self, doc, table_elem):
        """转换 HTML 表格为 docx 表格"""
        try:
            from docx.shared import Inches
        except ImportError:
            return

        rows = table_elem.find_all('tr')
        if not rows:
            return

        # 计算最大列数
        max_cols = 0
        for row in rows:
            cells = row.find_all(['td', 'th'])
            max_cols = max(max_cols, len(cells))

        if max_cols == 0:
            return

        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    table.cell(i, j).text = cell.get_text(strip=True)

    def file_extension(self) -> str:
        return ".docx"
