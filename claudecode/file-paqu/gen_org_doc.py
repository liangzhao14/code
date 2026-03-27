"""生成项目组织分工Word文档 - 黑白脑图，左→右树形展开"""
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# ── 字体 ──
font_path = 'C:/Windows/Fonts/msyh.ttc'
fp_center = FontProperties(fname=font_path, size=13, weight='bold')
fp_role = FontProperties(fname=font_path, size=9.5, weight='bold')
fp_detail = FontProperties(fname=font_path, size=7.5)

# ── 数据 ──
roles = [
    ('项目负责人',        '1人',    '统筹决策 / 资源协调'),
    ('业务负责人（甲方）', '1人',   'KRT数据 / 内容审核'),
    ('数据工程师',        '2人',    '采集 / 清洗 / 脱敏'),
    ('标注负责人',        '1人',    '试标 / 口径统一 / 考核'),
    ('标注业务人员',      '6~10人', '批量标注 / 核心产能'),
    ('质检人员',          '2人',    '三级质检 / 整改台账'),
    ('核电领域专家终审',   '1人',   '终审裁定 / 标准修订'),
    ('技术支持人员',      '1人',    '平台工具 / 环境保障'),
]

n = len(roles)

# ── 布局参数 ──
fig_w, fig_h = 13, 7
root_x = 1.2          # 根节点中心x
root_y = fig_h / 2    # 根节点中心y
root_w, root_h = 2.2, 1.2

node_x = 6.8          # 二级节点中心x
node_w, node_h = 3.6, 0.6
y_margin = 0.35
total_h = n * node_h + (n - 1) * y_margin
y_start = root_y + total_h / 2 - node_h / 2  # 顶部节点y

trunk_x = 3.8          # 主干线x位置

# ── 绘图 ──
fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=200)
ax.set_xlim(0, fig_w)
ax.set_ylim(0, fig_h)
ax.set_aspect('equal')
ax.axis('off')
fig.patch.set_facecolor('white')

# 根节点
rect = patches.FancyBboxPatch(
    (root_x - root_w/2, root_y - root_h/2), root_w, root_h,
    boxstyle="round,pad=0.15", linewidth=2, edgecolor='black', facecolor='white'
)
ax.add_patch(rect)
ax.text(root_x, root_y + 0.15, '项目组织架构', fontproperties=fp_center,
        ha='center', va='center', color='black')
ax.text(root_x, root_y - 0.25, '14~18 人', fontproperties=FontProperties(fname=font_path, size=9),
        ha='center', va='center', color='#444444')

# 根节点 → 主干线（水平）
ax.plot([root_x + root_w/2, trunk_x], [root_y, root_y],
        color='black', linewidth=1.5, solid_capstyle='round')

# 主干线（垂直）
top_y = y_start
bot_y = y_start - (n - 1) * (node_h + y_margin)
ax.plot([trunk_x, trunk_x], [bot_y, top_y],
        color='black', linewidth=1.5, solid_capstyle='round')

# 二级节点
for i, (role, count, duty) in enumerate(roles):
    yy = y_start - i * (node_h + y_margin)

    # 水平分支线
    ax.plot([trunk_x, node_x - node_w/2], [yy, yy],
            color='black', linewidth=1, solid_capstyle='round')

    # 节点框
    rect = patches.FancyBboxPatch(
        (node_x - node_w/2, yy - node_h/2), node_w, node_h,
        boxstyle="round,pad=0.08", linewidth=1.2, edgecolor='black', facecolor='white'
    )
    ax.add_patch(rect)

    # 角色名 + 人数 + 职责（单行排列）
    text_left = node_x - node_w/2 + 0.15
    ax.text(text_left, yy + 0.02, f'{role}  [{count}]', fontproperties=fp_role,
            ha='left', va='center', color='black')
    ax.text(node_x + node_w/2 - 0.15, yy + 0.02, duty, fontproperties=fp_detail,
            ha='right', va='center', color='#555555')

plt.tight_layout(pad=0.3)
img_path = os.path.join(os.path.dirname(__file__), '_org_chart.png')
plt.savefig(img_path, bbox_inches='tight', facecolor='white', edgecolor='none')
plt.close()
print(f'图片已生成: {img_path}')

# ── Word文档 ──
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2)
sec.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run('项目组织分工与实施保障')
run.bold = True
run.font.size = Pt(20)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(20)
run = sub.add_run('核电领域数据治理项目')
run.font.size = Pt(10.5)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 插入脑图
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(img_path, width=Cm(17))

out_path = os.path.join(os.path.dirname(__file__), '项目组织分工与实施保障.docx')
doc.save(out_path)
print(f'Word文档已生成: {out_path}')
